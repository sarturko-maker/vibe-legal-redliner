"""
Tests for doc_analyser.py — Document Structure Analyser.

Run: python3 test_doc_analyser.py
From: vibe-legal-extension/python/
"""

import sys
import os
from io import BytesIO
from zipfile import ZipFile

sys.path.insert(0, '.')

from lxml import etree
from docx import Document
from docx.oxml.ns import qn

from doc_analyser import (
    build_context_header,
    analyse_numbering,
    analyse_styles,
    build_paragraph_map,
)

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


# ---------------------------------------------------------------------------
# Helpers — build minimal .docx bytes for testing
# ---------------------------------------------------------------------------

def _doc_to_bytes(doc):
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_xml_parts(docx_bytes):
    """Extract parsed XML parts needed by analyser functions."""
    with ZipFile(BytesIO(docx_bytes)) as zf:
        doc_xml = etree.fromstring(zf.read('word/document.xml'))
        styles_xml = etree.fromstring(zf.read('word/styles.xml')) if 'word/styles.xml' in zf.namelist() else None
        numbering_xml = etree.fromstring(zf.read('word/numbering.xml')) if 'word/numbering.xml' in zf.namelist() else None
    body = doc_xml.find(f'{W}body')
    paragraphs = body.findall(f'{W}p')
    return paragraphs, numbering_xml, styles_xml


def _make_auto_numbered_doc():
    """Build a doc with auto-numbered list paragraphs."""
    doc = Document()
    doc.add_heading('DEFINITIONS', level=1)

    # Add numbered list items using ListParagraph-like approach
    # python-docx doesn't have direct list support, so we add numPr manually
    for i, text in enumerate(['In this Agreement:', 'Confidential Information means...', 'Purpose means...']):
        p = doc.add_paragraph(text)
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(p._element, qn('w:pPr'))
            p._element.insert(0, pPr)
        numPr = etree.SubElement(pPr, qn('w:numPr'))
        ilvl = etree.SubElement(numPr, qn('w:ilvl'))
        ilvl.set(qn('w:val'), '0')
        numId = etree.SubElement(numPr, qn('w:numId'))
        numId.set(qn('w:val'), '1')

    doc.add_heading('OBLIGATIONS', level=1)
    p = doc.add_paragraph('The Receiving Party shall keep all Confidential Information strictly confidential.')
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(p._element, qn('w:pPr'))
        p._element.insert(0, pPr)
    numPr = etree.SubElement(pPr, qn('w:numPr'))
    ilvl = etree.SubElement(numPr, qn('w:ilvl'))
    ilvl.set(qn('w:val'), '0')
    numId = etree.SubElement(numPr, qn('w:numId'))
    numId.set(qn('w:val'), '1')

    return _doc_to_bytes(doc)


def _make_manual_numbered_doc():
    """Build a doc with manual numbering (literal numbers in text)."""
    doc = Document()
    doc.add_heading('DEFINITIONS', level=1)
    doc.add_paragraph('1.\tIn this Agreement:')
    doc.add_paragraph('(a)\tConfidential Information means...')
    doc.add_paragraph('(b)\tPurpose means...')
    doc.add_heading('OBLIGATIONS', level=1)
    doc.add_paragraph('2.\tThe Receiving Party shall keep all information confidential.')
    doc.add_heading('GOVERNING LAW', level=1)
    doc.add_paragraph('3.\tThis Agreement shall be governed by English law.')
    return _doc_to_bytes(doc)


def _make_no_numbering_doc():
    """Build a doc with no numbering at all."""
    doc = Document()
    doc.add_paragraph('NON-DISCLOSURE AGREEMENT')
    doc.add_paragraph('This Agreement is entered into between the parties.')
    doc.add_paragraph('All information shall be kept confidential.')
    doc.add_paragraph('This Agreement is governed by English law.')
    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# Test 1: Auto-numbering detection
# ---------------------------------------------------------------------------

def test_auto_numbering_detection():
    """Documents with w:numPr should be detected as automatic numbering."""
    docx_bytes = _make_auto_numbered_doc()
    paragraphs, numbering_xml, styles_xml = _extract_xml_parts(docx_bytes)
    result = analyse_numbering(paragraphs, numbering_xml, styles_xml)

    assert result['scheme'] == 'automatic', f"Expected 'automatic', got '{result['scheme']}'"
    print("PASS: test_auto_numbering_detection")


# ---------------------------------------------------------------------------
# Test 2: Manual numbering detection
# ---------------------------------------------------------------------------

def test_manual_numbering_detection():
    """Documents with literal numbers should be detected as manual numbering."""
    docx_bytes = _make_manual_numbered_doc()
    paragraphs, numbering_xml, styles_xml = _extract_xml_parts(docx_bytes)
    result = analyse_numbering(paragraphs, numbering_xml, styles_xml)

    assert result['scheme'] == 'manual', f"Expected 'manual', got '{result['scheme']}'"
    print("PASS: test_manual_numbering_detection")


# ---------------------------------------------------------------------------
# Test 3: No numbering detection
# ---------------------------------------------------------------------------

def test_no_numbering_detection():
    """Documents with no numbering should be detected as 'none'."""
    docx_bytes = _make_no_numbering_doc()
    paragraphs, numbering_xml, styles_xml = _extract_xml_parts(docx_bytes)
    result = analyse_numbering(paragraphs, numbering_xml, styles_xml)

    assert result['scheme'] == 'none', f"Expected 'none', got '{result['scheme']}'"
    print("PASS: test_no_numbering_detection")


# ---------------------------------------------------------------------------
# Test 4: Paragraph map generation
# ---------------------------------------------------------------------------

def test_paragraph_map_generation():
    """Paragraph map should include style and content for non-empty paragraphs."""
    docx_bytes = _make_manual_numbered_doc()
    paragraphs, numbering_xml, _ = _extract_xml_parts(docx_bytes)
    result = build_paragraph_map(paragraphs, numbering_xml)

    assert 'DEFINITIONS' in result, f"Should contain heading text, got: {result[:200]}"
    assert '[' in result, f"Should contain paragraph indices"
    assert 'Heading1' in result or 'Heading 1' in result or 'heading' in result.lower(), \
        f"Should reference heading style"
    # Check manual numbers appear in content
    assert '1.' in result, f"Should show manual clause number in content"
    print("PASS: test_paragraph_map_generation")


# ---------------------------------------------------------------------------
# Test 5: Style analysis
# ---------------------------------------------------------------------------

def test_style_analysis():
    """Style analysis should detect heading and body styles."""
    docx_bytes = _make_auto_numbered_doc()
    paragraphs, _, styles_xml = _extract_xml_parts(docx_bytes)
    numbering_xml = None
    with ZipFile(BytesIO(docx_bytes)) as zf:
        if 'word/numbering.xml' in zf.namelist():
            numbering_xml = etree.fromstring(zf.read('word/numbering.xml'))

    result = analyse_styles(paragraphs, styles_xml, numbering_xml)

    assert len(result['heading_styles']) > 0, f"Should detect heading styles, got: {result['heading_styles']}"
    print("PASS: test_style_analysis")


# ---------------------------------------------------------------------------
# Test 6: Context header format
# ---------------------------------------------------------------------------

def test_context_header_format():
    """build_context_header should return a well-structured string."""
    docx_bytes = _make_auto_numbered_doc()
    header = build_context_header(docx_bytes)

    assert 'DOCUMENT STRUCTURE ANALYSIS:' in header, "Should contain structure analysis section"
    assert 'NUMBERING RULES FOR THIS DOCUMENT:' in header, "Should contain numbering rules"
    assert 'PARAGRAPH MAP:' in header, "Should contain paragraph map"
    assert 'automatic' in header.lower(), "Should mention automatic numbering"
    assert 'Do NOT include clause numbers' in header, "Should have auto-numbering instruction"
    print("PASS: test_context_header_format")


# ---------------------------------------------------------------------------
# Test 7: Manual numbering context header
# ---------------------------------------------------------------------------

def test_manual_numbering_context_header():
    """Manual numbering doc should get sub-numbering instructions."""
    docx_bytes = _make_manual_numbered_doc()
    header = build_context_header(docx_bytes)

    assert 'manual' in header.lower(), f"Should mention manual numbering"
    assert 'sub-numbering' in header.lower(), f"Should mention sub-numbering"
    assert '4A.' in header, "Should give sub-numbering example"
    print("PASS: test_manual_numbering_context_header")


# ---------------------------------------------------------------------------
# Test 8: Auto-numbered paragraph map shows AUTO-NUMBERED tag
# ---------------------------------------------------------------------------

def test_auto_numbered_tags_in_map():
    """Paragraphs with numPr should be tagged AUTO-NUMBERED in the map."""
    docx_bytes = _make_auto_numbered_doc()
    paragraphs, numbering_xml, _ = _extract_xml_parts(docx_bytes)
    result = build_paragraph_map(paragraphs, numbering_xml)

    assert 'AUTO-NUMBERED' in result, f"Should tag auto-numbered paragraphs, got: {result[:300]}"
    print("PASS: test_auto_numbered_tags_in_map")


# ---------------------------------------------------------------------------
# Test 9: Integration — real NDA files (if available)
# ---------------------------------------------------------------------------

def test_integration_real_ndas():
    """Run build_context_header on real test NDAs if available."""
    test_dir = os.path.expanduser('~/VibeLegalServer/Uploads/')
    if not os.path.exists(test_dir):
        print("SKIP: test_integration_real_ndas (no test dir)")
        return

    test_files = {}
    for fname in os.listdir(test_dir):
        if 'Test-NDA.docx' in fname and 'Manual' not in fname:
            test_files.setdefault('auto', fname)
        elif 'Manual-Numbering' in fname:
            test_files.setdefault('manual', fname)
        elif 'NDA-Simple-Numbered' in fname:
            test_files.setdefault('simple', fname)
        elif 'NDA-Bullet-Points' in fname:
            test_files.setdefault('bullet', fname)
        elif 'NDA-Word-Lists' in fname:
            test_files.setdefault('lists', fname)

    if not test_files:
        print("SKIP: test_integration_real_ndas (no test files found)")
        return

    for label, fname in test_files.items():
        path = os.path.join(test_dir, fname)
        with open(path, 'rb') as f:
            data = f.read()

        header = build_context_header(data)

        # Basic sanity checks
        assert 'DOCUMENT STRUCTURE ANALYSIS:' in header, f"{label}: missing structure analysis"
        assert 'PARAGRAPH MAP:' in header, f"{label}: missing paragraph map"
        assert 'NUMBERING RULES' in header, f"{label}: missing numbering rules"

        # Type-specific checks
        if label == 'auto':
            assert 'automatic' in header.lower(), f"{label}: should detect automatic numbering"
        elif label == 'manual':
            assert 'manual' in header.lower(), f"{label}: should detect manual numbering"
        elif label == 'simple':
            assert 'manual' in header.lower(), f"{label}: should detect manual numbering"

    print(f"PASS: test_integration_real_ndas ({len(test_files)} files tested: {', '.join(test_files.keys())})")


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    tests = [
        test_auto_numbering_detection,
        test_manual_numbering_detection,
        test_no_numbering_detection,
        test_paragraph_map_generation,
        test_style_analysis,
        test_context_header_format,
        test_manual_numbering_context_header,
        test_auto_numbered_tags_in_map,
        test_integration_real_ndas,
    ]

    passed = 0
    failed = 0
    for t in tests:
        try:
            t()
            passed += 1
        except Exception as e:
            print(f"FAIL: {t.__name__} — {e}")
            failed += 1

    print(f"\n{'=' * 50}")
    print(f"Results: {passed} passed, {failed} failed out of {len(tests)} tests")
    if failed > 0:
        sys.exit(1)
    else:
        print("All tests passed!")
