"""
Tests for styler.py — Styler post-processor.

Run: python3 test_styler.py
From: vibe-legal-extension/python/
"""

import sys
import copy
from io import BytesIO

sys.path.insert(0, '.')

from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

from styler import Styler, StylerResult

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
AUTHOR = "Vibe Legal"


# ---------------------------------------------------------------------------
# Helpers — build minimal .docx bytes with track changes
# ---------------------------------------------------------------------------

def _doc_to_bytes(doc):
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_vl_insertion(doc, text, bold=False, indent_twips=None, space_after_twips=None, numPr=False):
    """
    Add a paragraph wrapped in w:ins author="Vibe Legal".
    This simulates what Adeu's RedlineEngine produces.
    """
    p = doc.add_paragraph()
    p_elem = p._element

    # Build run
    run_elem = etree.SubElement(p_elem, qn('w:r'))
    if bold:
        rPr = etree.SubElement(run_elem, qn('w:rPr'))
        etree.SubElement(rPr, qn('w:b'))
    t = etree.SubElement(run_elem, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')

    # Wrap all child elements in w:ins
    ins = etree.Element(qn('w:ins'))
    ins.set(qn('w:id'), str(id(p_elem) % 999999))
    ins.set(qn('w:author'), AUTHOR)
    ins.set(qn('w:date'), '2025-01-01T00:00:00Z')

    # Move run into ins
    p_elem.remove(run_elem)
    ins.append(run_elem)
    p_elem.append(ins)

    # Paragraph properties
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(p_elem, qn('w:pPr'))
        p_elem.insert(0, pPr)

    if indent_twips is not None:
        ind = etree.SubElement(pPr, qn('w:ind'))
        ind.set(qn('w:left'), str(indent_twips))

    if space_after_twips is not None:
        spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:after'), str(space_after_twips))

    if numPr:
        nPr = etree.SubElement(pPr, qn('w:numPr'))
        ilvl = etree.SubElement(nPr, qn('w:ilvl'))
        ilvl.set(qn('w:val'), '0')
        numId = etree.SubElement(nPr, qn('w:numId'))
        numId.set(qn('w:val'), '1')

    return p


def _add_original_paragraph(doc, text, bold=False, indent_twips=None, space_after_twips=None, is_header=False):
    """Add a regular (non-VL) paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True

    if indent_twips is not None:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(p._element, qn('w:pPr'))
            p._element.insert(0, pPr)
        ind = etree.SubElement(pPr, qn('w:ind'))
        ind.set(qn('w:left'), str(indent_twips))

    if space_after_twips is not None:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(p._element, qn('w:pPr'))
            p._element.insert(0, pPr)
        spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:after'), str(space_after_twips))

    return p


# ---------------------------------------------------------------------------
# Test 1: Section header gets bold applied
# ---------------------------------------------------------------------------

def test_section_header_bold():
    """VL insertion that looks like a section header should get bold formatting."""
    # Build original doc with a bold section header (for reference detection)
    orig = Document()
    _add_original_paragraph(orig, "DEFINITIONS:", bold=True)
    _add_original_paragraph(orig, "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.")
    orig_bytes = _doc_to_bytes(orig)

    # Build redlined doc with a VL-inserted section header that's NOT bold
    redlined = Document()
    _add_original_paragraph(redlined, "DEFINITIONS:", bold=True)
    _add_original_paragraph(redlined, "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.")
    _add_vl_insertion(redlined, "OBLIGATIONS:")  # No bold, no numPr

    redlined_bytes = _doc_to_bytes(redlined)

    # Detect reference from original
    ref_styler = Styler(orig_bytes, author=AUTHOR)
    ref_formats = ref_styler.detect_reference_formats()

    # Run styler on redlined doc
    styler = Styler(redlined_bytes, author=AUTHOR, original_reference=ref_formats)
    result = styler.run()

    assert result.fix_count > 0, f"Expected fixes, got {result.fix_count}"
    assert any("BOLD" in f or "bold" in f.lower() for f in result.fixes_applied), \
        f"Expected bold fix, got: {result.fixes_applied}"
    print("PASS: test_section_header_bold")


# ---------------------------------------------------------------------------
# Test 2: Inline title bold
# ---------------------------------------------------------------------------

def test_inline_title_bold():
    """VL-inserted numbered clause should get bold applied to inline title."""
    # Original has numbered clause with bold inline title
    orig = Document()
    p = orig.add_paragraph()
    run1 = p.add_run("1. Purpose.")
    run1.bold = True
    p.add_run(" The Parties wish to exchange confidential information.")
    orig_bytes = _doc_to_bytes(orig)

    # Redlined doc: VL inserted a numbered clause WITHOUT bold
    redlined = Document()
    # Keep original clause
    p2 = redlined.add_paragraph()
    r1 = p2.add_run("1. Purpose.")
    r1.bold = True
    p2.add_run(" The Parties wish to exchange confidential information.")
    # VL inserts a new clause (not bold)
    _add_vl_insertion(redlined, "2. Exclusions. The following shall not be considered Confidential Information.")

    redlined_bytes = _doc_to_bytes(redlined)

    ref_styler = Styler(orig_bytes, author=AUTHOR)
    ref_formats = ref_styler.detect_reference_formats()
    assert ref_formats['inline_title']['has_pattern'], "Should detect inline title pattern"
    assert ref_formats['inline_title']['title_is_bold'], "Should detect title is bold"

    styler = Styler(redlined_bytes, author=AUTHOR, original_reference=ref_formats)
    result = styler.run()

    assert result.fix_count > 0, f"Expected fixes, got {result.fix_count}"
    assert any("bold" in f.lower() for f in result.fixes_applied), \
        f"Expected bold fix, got: {result.fixes_applied}"
    print("PASS: test_inline_title_bold")


# ---------------------------------------------------------------------------
# Test 3: Body indentation
# ---------------------------------------------------------------------------

def test_body_indentation():
    """VL insertion body paragraphs should inherit indent from original."""
    orig = Document()
    _add_original_paragraph(orig, "DEFINITIONS:", bold=True)
    _add_original_paragraph(
        orig,
        "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.",
        indent_twips=720
    )
    orig_bytes = _doc_to_bytes(orig)

    redlined = Document()
    _add_original_paragraph(redlined, "DEFINITIONS:", bold=True)
    _add_original_paragraph(
        redlined,
        "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.",
        indent_twips=720
    )
    # VL inserts body paragraph WITHOUT indent
    _add_vl_insertion(
        redlined,
        "The Receiving Party shall protect all Confidential Information using reasonable measures for the entire term."
    )
    redlined_bytes = _doc_to_bytes(redlined)

    ref_styler = Styler(orig_bytes, author=AUTHOR)
    ref_formats = ref_styler.detect_reference_formats()
    assert ref_formats['body_indent']['left_indent'] == 720, \
        f"Should detect indent=720, got {ref_formats['body_indent']['left_indent']}"

    styler = Styler(redlined_bytes, author=AUTHOR, original_reference=ref_formats)
    result = styler.run()

    assert result.fix_count > 0, f"Expected indent fixes, got {result.fix_count}"
    assert any("indent" in f.lower() for f in result.fixes_applied), \
        f"Expected indent fix, got: {result.fixes_applied}"
    print("PASS: test_body_indentation")


# ---------------------------------------------------------------------------
# Test 4: Spacing applied
# ---------------------------------------------------------------------------

def test_spacing_applied():
    """VL insertion paragraphs should inherit space_after from original."""
    orig = Document()
    _add_original_paragraph(orig, "DEFINITIONS:", bold=True, space_after_twips=200)
    _add_original_paragraph(
        orig,
        "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.",
        space_after_twips=200
    )
    orig_bytes = _doc_to_bytes(orig)

    redlined = Document()
    _add_original_paragraph(redlined, "DEFINITIONS:", bold=True, space_after_twips=200)
    _add_original_paragraph(
        redlined,
        "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.",
        space_after_twips=200
    )
    # VL inserts without spacing
    _add_vl_insertion(redlined, "OBLIGATIONS:")
    _add_vl_insertion(
        redlined,
        "The Receiving Party shall protect all Confidential Information using reasonable measures for the entire term."
    )
    redlined_bytes = _doc_to_bytes(redlined)

    ref_styler = Styler(orig_bytes, author=AUTHOR)
    ref_formats = ref_styler.detect_reference_formats()
    assert ref_formats['spacing']['space_after'] == 200, \
        f"Should detect space_after=200, got {ref_formats['spacing']['space_after']}"

    styler = Styler(redlined_bytes, author=AUTHOR, original_reference=ref_formats)
    result = styler.run()

    assert result.fix_count > 0, f"Expected spacing fixes, got {result.fix_count}"
    assert any("space_after" in f for f in result.fixes_applied), \
        f"Expected space_after fix, got: {result.fixes_applied}"
    print("PASS: test_spacing_applied")


# ---------------------------------------------------------------------------
# Test 5: Original paragraphs not modified, issues in warnings
# ---------------------------------------------------------------------------

def test_original_not_modified():
    """
    Styler must NOT modify original (non-VL) paragraphs.
    Issues with originals should appear as warnings.
    """
    doc = Document()
    # Original section header with numPr (wrong formatting — should be flagged)
    p = _add_original_paragraph(doc, "OBLIGATIONS:")
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(p._element, qn('w:pPr'))
        p._element.insert(0, pPr)
    nPr = etree.SubElement(pPr, qn('w:numPr'))
    ilvl = etree.SubElement(nPr, qn('w:ilvl'))
    ilvl.set(qn('w:val'), '0')
    numId = etree.SubElement(nPr, qn('w:numId'))
    numId.set(qn('w:val'), '1')

    _add_original_paragraph(doc, "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.")

    doc_bytes = _doc_to_bytes(doc)

    styler = Styler(doc_bytes, author=AUTHOR)
    result = styler.run()

    # Should have warnings about the original paragraph
    assert result.has_warnings, f"Expected warnings about original paragraph, got none"
    assert any("original paragraph" in w for w in result.warnings), \
        f"Expected 'original paragraph' in warnings, got: {result.warnings}"
    # Should NOT have fixes (nothing to fix on originals)
    assert result.fix_count == 0, f"Expected 0 fixes on originals, got {result.fix_count}"
    print("PASS: test_original_not_modified")


# ---------------------------------------------------------------------------
# Test 6: Idempotent
# ---------------------------------------------------------------------------

def test_idempotent():
    """Running Styler twice should produce the same result."""
    orig = Document()
    _add_original_paragraph(orig, "DEFINITIONS:", bold=True, space_after_twips=200)
    _add_original_paragraph(
        orig,
        "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.",
        indent_twips=720, space_after_twips=200
    )
    orig_bytes = _doc_to_bytes(orig)

    redlined = Document()
    _add_original_paragraph(redlined, "DEFINITIONS:", bold=True, space_after_twips=200)
    _add_original_paragraph(
        redlined,
        "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.",
        indent_twips=720, space_after_twips=200
    )
    _add_vl_insertion(redlined, "OBLIGATIONS:")
    _add_vl_insertion(
        redlined,
        "The Receiving Party shall protect all Confidential Information using reasonable measures for the entire term."
    )
    redlined_bytes = _doc_to_bytes(redlined)

    ref_styler = Styler(orig_bytes, author=AUTHOR)
    ref_formats = ref_styler.detect_reference_formats()

    # First run
    s1 = Styler(redlined_bytes, author=AUTHOR, original_reference=ref_formats)
    r1 = s1.run()
    assert r1.fix_count > 0, "First run should apply fixes"

    # Second run on already-fixed output
    s2 = Styler(r1.modified_bytes, author=AUTHOR, original_reference=ref_formats)
    r2 = s2.run()
    assert r2.fix_count == 0, f"Second run should find 0 fixes, got {r2.fix_count}: {r2.fixes_applied}"

    print("PASS: test_idempotent")


# ---------------------------------------------------------------------------
# Test 7: No fixes needed
# ---------------------------------------------------------------------------

def test_no_fixes_needed():
    """Document with correct formatting should produce 0 fixes."""
    doc = Document()
    _add_original_paragraph(doc, "DEFINITIONS:", bold=True, space_after_twips=200)
    _add_original_paragraph(
        doc,
        "This Agreement sets forth the following terms and conditions between the parties involved in the transaction.",
        indent_twips=720, space_after_twips=200
    )
    # VL insertion that already has correct formatting
    _add_vl_insertion(doc, "OBLIGATIONS:", bold=True, space_after_twips=200)
    _add_vl_insertion(
        doc,
        "The Receiving Party shall protect all Confidential Information using reasonable measures for the entire term.",
        indent_twips=720, space_after_twips=200
    )
    doc_bytes = _doc_to_bytes(doc)

    styler = Styler(doc_bytes, author=AUTHOR)
    result = styler.run()

    assert result.fix_count == 0, f"Expected 0 fixes, got {result.fix_count}: {result.fixes_applied}"
    print("PASS: test_no_fixes_needed")


# ---------------------------------------------------------------------------
# Test 8: StylerResult dataclass
# ---------------------------------------------------------------------------

def test_styler_result_dataclass():
    """Verify StylerResult properties."""
    r = StylerResult()
    assert r.fix_count == 0
    assert not r.has_warnings
    assert r.modified_bytes == b''

    r.fixes_applied.append("fix1")
    r.warnings.append("warn1")
    assert r.fix_count == 1
    assert r.has_warnings
    print("PASS: test_styler_result_dataclass")


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    tests = [
        test_styler_result_dataclass,
        test_section_header_bold,
        test_inline_title_bold,
        test_body_indentation,
        test_spacing_applied,
        test_original_not_modified,
        test_idempotent,
        test_no_fixes_needed,
    ]

    passed = 0
    failed = 0
    for t in tests:
        try:
            t()
            passed += 1
        except Exception as e:
            print(f"FAIL: {t.__name__} — {e}")
            failed += 1

    print(f"\n{'=' * 50}")
    print(f"Results: {passed} passed, {failed} failed out of {len(tests)} tests")
    if failed > 0:
        sys.exit(1)
    else:
        print("All tests passed!")
